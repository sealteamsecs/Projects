# Import docx NOT python-docx
import docx
from docx.shared import Pt
from docx.shared import Inches
from bs4 import BeautifulSoup
from markdown import markdown
import os
from collections import defaultdict
from pycti import OpenCTIApiClient

# Configure the OpenCTI client
api_url = os.getenv('OPENCTI_API_URL', '<OPENCTI_URL>')
api_token = os.getenv('OPENCTI_API_TOKEN', '<API_KEY>')

# Initialize the OpenCTI client
opencti_api_client = OpenCTIApiClient(api_url, api_token)

# Define the malware profile name
malware_profile_name = "<MALWARE NAME (e.g. 'QakBot')"

# Fetch the malware profile
malware = opencti_api_client.malware.read(filters=[{"key": "name", "values": [malware_profile_name]}])

# Get Threat Actor profile (Intrusion Sets)
#malware = opencti_api_client.intrusion_set.read(filters=[{"key": "name", "values": [malware_profile_name]}])

# Create a Word Document
doc = docx.Document()
 

###### Create the Header Template ######

# Choosing the top most section of the page
section = doc.sections[0]

# Selecting the header
header = section.header




# Creates a paragraph in the header section 
header_para = header.paragraphs[0]
run = header_para.add_run()

# Add Logo to top-right of the header
run.add_picture("logo.png", width=Inches(1.25))

# Add the Traffic Light Protocol
header_para = header.paragraphs[0]
run = header_para.add_run()



###### End Header Template Section ######




if not malware:
    print(f"Malware profile '{malware_profile_name}' not found.")
else:
    
    ###### Add the Title Section ######

    # Add a Title to the document

    p = doc.add_heading('MAIN HEADING',2)
    p.bold = True
    p.alignment = 1

    # Pulls Name from OpenCTI Json object
    p = doc.add_paragraph(malware['name'])
    p.bold = True
    p.alignment = 1

    ###### end the Title Section ######

    
    ###### Add the Summary Section ######

    description = malware['description']

    htmldescription = markdown(description)
    html = markdown(htmldescription)
    text = ''.join(BeautifulSoup(htmldescription, "lxml").findAll(string=True))

    para = doc.add_paragraph(text)

    ###### End the Summary Section ######   

    ###### Add the Indicators Section ######
    p = doc.add_heading('Indicators', 2)
    
    
    
    # Fetch the 20 most recent indicators related to the malware profile
    indicators = opencti_api_client.indicator.list(
        filters=[{"key": "indicates", "values": [malware["id"]]}],
        orderBy="created_at",
        orderMode="desc",
        first=20
    )
    # Group indicators by pattern type
    grouped_indicators = defaultdict(list)
    for indicator in indicators:
        grouped_indicators[indicator['x_opencti_main_observable_type']].append(indicator)

    # Print the grouped indicators
    if not grouped_indicators:
        print(f"No indicators found for malware profile '{malware_profile_name}'.")
    else:
        for pattern_type, indicator_list in grouped_indicators.items():
            for index, indicator in enumerate(indicator_list):
                paragraph = doc.add_paragraph(indicator['name'])
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(5)                
                #print(f"{index + 1}. Indicator '{indicator['name']}': {indicator['pattern']}")

    ###### End the Indicator Section ###### 

    ###### Add the Victims Section ######
    p = doc.add_heading('Victims', 2)

    # Fetch the victims related to the malware profile
    victims = opencti_api_client.stix_core_relationship.list(
        fromId=malware['id'],
        fromTypes=['Malware'],
        toTypes=['Organization', 'Sector', 'Region', 'Country'],
        relationship_type='targets',
        first=20
    )

    if not victims:
        print(f"No victims found for malware profile '{malware_profile_name}'.")
    else:
        for index, victim in enumerate(victims):
            victim_entity = victim['to']
            paragraph = doc.add_paragraph(f"{index + 1}. Victim '{victim_entity['name']}': {victim_entity['entity_type']}")
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(5)  
    ###### End the Victims Section ######    
# Now save the document to a location
doc.save('malware_profile.docx')
